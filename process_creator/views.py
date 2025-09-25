from django.shortcuts import get_object_or_404, redirect, render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.db import transaction
from django.contrib.auth.decorators import login_required
from rbac.decorators import require_app_access
from django.db import models
from django.template.loader import render_to_string
from django.conf import settings
from django.utils import timezone
from .models import Module, Process, Step, StepImage, StepLink, StepFile, AIInteraction, ProcessTemplate, TemplateStep, Job, JobStep, JobSubtask, JobStepImage
from .conf import JOB_LABEL
from .services.templates import sync_process_to_template
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import openai
import logging

logger = logging.getLogger(__name__)
import re
from decimal import Decimal
from io import BytesIO
from urllib.parse import quote
import subprocess
import tempfile

# Optional PDF rendering (to embed PDF pages as images in Word)
def _render_pdf_pages_to_images(pdf_path: str):
    images = []
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(pdf_path)
        try:
            for page in doc:
                pix = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
                img_bytes = pix.tobytes("png")
                bio = BytesIO(img_bytes)
                images.append(bio)
        finally:
            doc.close()
    except Exception:
        return []
    return images


def markdown_to_plain_text(text):
    """Convert basic Markdown formatting to plain text with proper formatting"""
    if not text:
        return ""
    
    # Convert headers
    text = re.sub(r'^### (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.+)$', r'\1', text, flags=re.MULTILINE)
    
    # Convert bold text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    
    # Convert italic text
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    
    # Convert bullet points
    text = re.sub(r'^\* (.+)$', r'• \1', text, flags=re.MULTILINE)
    text = re.sub(r'^- (.+)$', r'• \1', text, flags=re.MULTILINE)
    
    # Convert numbered lists
    text = re.sub(r'^(\d+)\. (.+)$', r'\1. \2', text, flags=re.MULTILINE)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


def add_markdown_to_word_doc(doc, text, level=1):
    """Add Markdown-formatted text to a Word document with proper formatting"""
    if not text:
        return
    
    lines = text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Handle headers
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=level)
            heading.paragraph_format.space_after = Inches(0.1)
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=level + 1)
            heading.paragraph_format.space_after = Inches(0.1)
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=level + 2)
            heading.paragraph_format.space_after = Inches(0.1)
        
        # Handle bullet points
        elif line.startswith('- ') or line.startswith('* '):
            # Collect all consecutive bullet points
            bullet_items = []
            while i < len(lines) and (lines[i].strip().startswith('- ') or lines[i].strip().startswith('* ')):
                bullet_text = lines[i].strip()[2:].strip()
                # Handle bold text in bullets
                bullet_text = re.sub(r'\*\*(.+?)\*\*', r'\1', bullet_text)
                bullet_items.append(bullet_text)
                i += 1
            i -= 1  # Back up one since we'll increment at the end
            
            # Add bullet list
            for item in bullet_items:
                p = doc.add_paragraph(item, style='List Bullet')
                p.paragraph_format.space_after = Inches(0.05)
                p.paragraph_format.line_spacing = 1.15
        
        # Handle numbered lists
        elif re.match(r'^\d+\. ', line):
            # Collect all consecutive numbered items
            numbered_items = []
            while i < len(lines) and re.match(r'^\d+\. ', lines[i].strip()):
                item_text = re.sub(r'^\d+\. ', '', lines[i].strip())
                # Handle bold text in numbered items
                item_text = re.sub(r'\*\*(.+?)\*\*', r'\1', item_text)
                numbered_items.append(item_text)
                i += 1
            i -= 1  # Back up one since we'll increment at the end
            
            # Add numbered list
            for item in numbered_items:
                p = doc.add_paragraph(item, style='List Number')
                p.paragraph_format.space_after = Inches(0.05)
                p.paragraph_format.line_spacing = 1.15
        
        # Handle regular paragraphs
        else:
            # Handle bold text
            paragraph_text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            
            # Check if this is part of a multi-line paragraph
            if i + 1 < len(lines) and lines[i + 1].strip() and not lines[i + 1].strip().startswith(('#', '-', '*')) and not re.match(r'^\d+\. ', lines[i + 1].strip()):
                # Collect the full paragraph
                full_paragraph = [line]
                i += 1
                while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith(('#', '-', '*')) and not re.match(r'^\d+\. ', lines[i].strip()):
                    full_paragraph.append(lines[i].strip())
                    i += 1
                i -= 1  # Back up one since we'll increment at the end
                paragraph_text = ' '.join(full_paragraph)
            
            p = doc.add_paragraph(paragraph_text)
            p.paragraph_format.space_after = Inches(0.1)
            p.paragraph_format.line_spacing = 1.15
        
        i += 1


@login_required
@require_app_access('process_creator', action='view')
def process_list(request):
    # Get all modules for the dropdown
    modules = Module.objects.all()
    
    # Get selected module from request
    selected_module_id = request.GET.get('module')
    selected_module = None
    
    if selected_module_id:
        try:
            selected_module = Module.objects.get(id=selected_module_id)
            processes = Process.objects.filter(module=selected_module).order_by('order')
        except Module.DoesNotExist:
            processes = Process.objects.all().order_by('order')
    else:
        processes = Process.objects.all().order_by('order')
    
    return render(request, "process_creator/list.html", {
        "processes": processes,
        "modules": modules,
        "selected_module": selected_module
    })


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def module_create(request):
    """Create a new Module via AJAX."""
    data = json.loads(request.body or '{}') if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip()
    if not name:
        return JsonResponse({"ok": False, "error": "Name is required"}, status=400)
    # Prevent duplicates (case-insensitive)
    existing = Module.objects.filter(name__iexact=name).first()
    if existing:
        return JsonResponse({"ok": True, "id": existing.id, "name": existing.name})
    module = Module.objects.create(name=name, description=description)
    return JsonResponse({"ok": True, "id": module.id, "name": module.name})


@login_required
@require_app_access('process_creator', action='view')
def module_manage(request):
    modules = Module.objects.all().order_by('name')
    return render(request, 'process_creator/modules.html', {"modules": modules})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def module_update(request, module_id: int):
    module = get_object_or_404(Module, id=module_id)
    data = json.loads(request.body or '{}') if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip()
    if not name:
        return JsonResponse({"ok": False, "error": "Name is required"}, status=400)
    # Prevent duplicate names
    if Module.objects.exclude(id=module.id).filter(name__iexact=name).exists():
        return JsonResponse({"ok": False, "error": "A module with this name already exists"}, status=400)
    module.name = name
    module.description = description
    module.save()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def module_delete(request, module_id: int):
    module = get_object_or_404(Module, id=module_id)
    # Optional: refuse delete if in use
    if Process.objects.filter(module=module).exists():
        return JsonResponse({"ok": False, "error": "Module is assigned to one or more processes"}, status=400)
    module.delete()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
def process_create(request):
    max_order = Process.objects.aggregate(models.Max("order")).get("order__max") or 0
    process = Process.objects.create(name="Untitled Process", order=max_order + 1)
    return redirect("process_creator:edit", pk=process.pk)


@login_required
@require_app_access('process_creator', action='edit')
def process_edit(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    modules = Module.objects.all()
    return render(request, "process_creator/edit.html", {"process": process, "modules": modules})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def process_update(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    name = request.POST.get('name')
    description = request.POST.get('description')
    notes = request.POST.get('notes')
    summary = request.POST.get('summary')
    summary_instructions = request.POST.get('summary_instructions')
    analysis = request.POST.get('analysis')
    analysis_instructions = request.POST.get('analysis_instructions')
    module_id = request.POST.get('module')
    
    if name is not None:
        if not name.strip():
            return JsonResponse({"ok": False, "error": "Name is required."}, status=400)
        process.name = name.strip()
    if description is not None:
        process.description = description
    if notes is not None:
        process.notes = notes
    if summary is not None:
        process.summary = summary
    if summary_instructions is not None:
        process.summary_instructions = summary_instructions
    if analysis is not None:
        process.analysis = analysis
    if analysis_instructions is not None:
        process.analysis_instructions = analysis_instructions
    if module_id is not None:
        if module_id == '':
            process.module = None
        else:
            try:
                process.module = Module.objects.get(id=module_id)
            except Module.DoesNotExist:
                pass  # Keep existing module if invalid ID provided
    
    process.save()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='delete')
@require_POST
def process_delete(request, pk: int):
    # Delete a process; support AJAX (JSON) and non-AJAX (redirect) callers
    try:
        process = Process.objects.get(pk=pk)
    except Process.DoesNotExist:
        # For AJAX callers, return ok so UI can remove the stale item
        if request.headers.get('x-requested-with') == 'XMLHttpRequest' or 'application/json' in request.headers.get('Accept',''):
            return JsonResponse({"ok": True, "removed": True})
        return redirect("process_creator:list")
    process.delete()
    if request.headers.get('x-requested-with') == 'XMLHttpRequest' or 'application/json' in request.headers.get('Accept',''):
        return JsonResponse({"ok": True})
    return redirect("process_creator:list")


@login_required
@require_app_access('process_creator', action='view')
def process_print(request, pk: int):
    try:
        process = Process.objects.get(pk=pk)
    except Process.DoesNotExist:
        return redirect("process_creator:list")
    focused_step_id = request.GET.get('focused_step')
    try:
        steps = process.steps.all() if not focused_step_id else process.steps.filter(id=int(focused_step_id))
    except (TypeError, ValueError):
        steps = process.steps.all()
    return render(request, "process_creator/print.html", {"process": process, "steps": steps})


@login_required
@require_app_access('process_creator', action='view')
def process_copy_prompt(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    lines = [f"# {process.name}"]
    if process.description:
        lines.append(process.description)
    lines.append("\n## Steps")
    for step in process.steps.all():
        title = step.title.strip() or "(Untitled Step)"
        details = step.details.strip()
        if details:
            lines.append(f"{step.order}. {title} — {details}")
        else:
            lines.append(f"{step.order}. {title}")
    if process.notes:
        lines.append("\n## Notes")
        lines.append(process.notes)
    text = "\n\n".join(lines)
    return HttpResponse(text, content_type="text/markdown; charset=utf-8")


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def processes_reorder(request):
    order_list = request.POST.getlist("order[]")
    with transaction.atomic():
        for index, process_id in enumerate(order_list, start=1):
            Process.objects.filter(id=process_id).update(order=index, updated_at=timezone.now())
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='view')
def process_print_all(request):
    # ids parameter optional; if provided, filter
    ids = request.GET.getlist('ids')
    processes = Process.objects.all()
    if ids:
        processes = processes.filter(id__in=ids)
    return render(request, "process_creator/print_all.html", {"processes": processes})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def steps_reorder(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    order_list = request.POST.getlist("order[]")
    with transaction.atomic():
        for index, step_id in enumerate(order_list, start=1):
            Step.objects.filter(process=process, id=step_id).update(order=index, updated_at=timezone.now())
        Process.objects.filter(pk=process.pk).update(updated_at=timezone.now())  # touch updated_at
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_add(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    max_order = process.steps.aggregate(models.Max("order")).get("order__max") or 0
    title = request.POST.get("title", "").strip()
    if not title:
        return JsonResponse({"ok": False, "error": "Title is required."}, status=400)
    step = Step.objects.create(
        process=process,
        order=max_order + 1,
        title=title,
        details=request.POST.get("details", ""),
    )
    return JsonResponse({"ok": True, "id": step.id, "order": step.order})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_insert(request, pk: int, step_id: int, direction: str):
    process = get_object_or_404(Process, pk=pk)
    ref = get_object_or_404(Step, pk=step_id, process=process)
    title = request.POST.get("title", "New Step").strip() or "New Step"
    if direction not in ("up", "down"):
        return JsonResponse({"ok": False, "error": "Invalid direction"}, status=400)
    with transaction.atomic():
        insert_order = ref.order if direction == "up" else ref.order + 1
        # shift steps >= insert_order by +1
        Step.objects.filter(process=process, order__gte=insert_order).update(order=models.F('order') + 1, updated_at=timezone.now())
        step = Step.objects.create(process=process, order=insert_order, title=title)
    return JsonResponse({"ok": True, "id": step.id, "order": step.order})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_update(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    title = request.POST.get("title")
    details = request.POST.get("details")
    if title is not None:
        if not title.strip():
            return JsonResponse({"ok": False, "error": "Title is required."}, status=400)
        step.title = title.strip()
    if details is not None:
        step.details = details
    step.save()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='delete')
@require_POST
def step_delete(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    step.delete()
    for index, s in enumerate(process.steps.order_by("order", "id"), start=1):
        if s.order != index:
            Step.objects.filter(pk=s.pk).update(order=index, updated_at=timezone.now())
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_image_upload(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    print(f"DEBUG: Uploading image to Process {process.id} ({process.name}), Step {step.id} ({step.title})")
    # Expect an image file under 'image' or a pasted blob as 'file'
    file = request.FILES.get('image') or request.FILES.get('file')
    if not file:
        return JsonResponse({"ok": False, "error": "No image provided"}, status=400)
    # Optional substep index (0-based) to associate image with a bullet line
    substep_index = request.POST.get('substep_index')
    try:
        substep_index = int(substep_index) if substep_index is not None and substep_index != '' else None
        if substep_index is not None and substep_index < 0:
            substep_index = None
    except (TypeError, ValueError):
        substep_index = None
    max_order = step.images.aggregate(models.Max('order')).get('order__max') or 0
    img = StepImage.objects.create(step=step, image=file, order=max_order + 1, substep_index=substep_index)
    return JsonResponse({"ok": True, "id": img.id, "url": img.image.url, "order": img.order, "substep_index": img.substep_index})


@login_required
@require_app_access('process_creator', action='delete')
@require_POST
def step_image_delete(request, pk: int, step_id: int, image_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    img = get_object_or_404(StepImage, pk=image_id, step=step)
    img.delete()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_image_update_substep(request, pk: int, step_id: int, image_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    img = get_object_or_404(StepImage, pk=image_id, step=step)
    
    data = json.loads(request.body or '{}') if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    substep_index = data.get('substep_index')
    
    if substep_index is not None:
        img.substep_index = substep_index
        img.save()
        return JsonResponse({"ok": True})
    else:
        return JsonResponse({"error": "substep_index is required"}, status=400)


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_images_clear_substeps(request, pk: int, step_id: int):
    """Clear substep_index for all images in a single step."""
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    updated = 0
    for img in step.images.all():
        if img.substep_index is not None:
            img.substep_index = None
            img.save(update_fields=["substep_index", "updated_at"])
            updated += 1
    return JsonResponse({"ok": True, "cleared": updated})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_link_add(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    data = json.loads(request.body or '{}') if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    title = (data.get('title') or '').strip()
    url = (data.get('url') or '').strip()
    if not title or not url:
        return JsonResponse({"ok": False, "error": "Title and URL are required"}, status=400)
    max_order = step.links.aggregate(models.Max('order')).get('order__max') or 0
    link = StepLink.objects.create(step=step, title=title, url=url, order=max_order + 1)
    return JsonResponse({"ok": True, "id": link.id, "title": link.title, "url": link.url})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_link_delete(request, pk: int, step_id: int, link_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    link = get_object_or_404(StepLink, pk=link_id, step=step)
    link.delete()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_file_upload(request, pk: int, step_id: int):
    try:
        process = get_object_or_404(Process, pk=pk)
        step = get_object_or_404(Step, pk=step_id, process=process)
        file = request.FILES.get('file')
        if not file:
            return JsonResponse({"ok": False, "error": "No file uploaded"}, status=200)
        name_lower = file.name.lower()
        max_order = step.files.aggregate(models.Max('order')).get('order__max') or 0

        # Direct PDF: save as-is
        if (file.content_type == 'application/pdf' or name_lower.endswith('.pdf')):
            sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
            return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": os.path.basename(sf.file.name)})

        # DWG / IDW: try local conversion to PDF
        is_dwg = name_lower.endswith('.dwg')
        is_idw = name_lower.endswith('.idw')
        if not (is_dwg or is_idw):
            return JsonResponse({"ok": False, "error": "Only PDF, DWG or IDW files are supported"}, status=200)

        # Persist the uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.name)[1]) as tmp_in:
            for chunk in file.chunks():
                tmp_in.write(chunk)
            tmp_in_path = tmp_in.name

        tmp_out_fd, tmp_out_path = tempfile.mkstemp(suffix='.pdf')
        os.close(tmp_out_fd)

        if is_dwg:
            exe = getattr(settings, 'ODA_CONVERTER_PDF', '')
            if not exe or not os.path.exists(exe):
                # Fallback: store the DWG as-is so it can be linked/downloaded
                sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
                return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": os.path.basename(sf.file.name)})
            cmd = [exe, tmp_in_path, tmp_out_path]
        else:
            exe = getattr(settings, 'INVENTOR_IDW_TO_PDF', '')
            if not exe or not os.path.exists(exe):
                sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
                return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": os.path.basename(sf.file.name)})
            cmd = [exe, tmp_in_path, tmp_out_path]

        completed = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, shell=False, timeout=120)
        if completed.returncode != 0 or not os.path.exists(tmp_out_path) or os.path.getsize(tmp_out_path) == 0:
            # On conversion failure, store original so user still gets a downloadable file
            try:
                os.unlink(tmp_out_path)
            except Exception:
                pass
            sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
            return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": os.path.basename(sf.file.name)})

        with open(tmp_out_path, 'rb') as f_out:
            from django.core.files.base import ContentFile
            pdf_bytes = f_out.read()
            pdf_name = os.path.splitext(file.name)[0] + '.pdf'
            sf = StepFile.objects.create(step=step, order=max_order + 1)
            sf.file.save(pdf_name, ContentFile(pdf_bytes), save=True)

        return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": os.path.basename(sf.file.name)})
    except Exception as e:
        return JsonResponse({"ok": False, "error": str(e)}, status=200)
    finally:
        try:
            if 'tmp_in_path' in locals() and os.path.exists(tmp_in_path):
                os.unlink(tmp_in_path)
        except Exception:
            pass
        try:
            if 'tmp_out_path' in locals() and os.path.exists(tmp_out_path):
                os.unlink(tmp_out_path)
        except Exception:
            pass


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_file_delete(request, pk: int, step_id: int, file_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    sf = get_object_or_404(StepFile, pk=file_id, step=step)
    sf.delete()
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def step_images_reorder(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    order_list = request.POST.getlist("order[]")
    with transaction.atomic():
        for index, img_id in enumerate(order_list, start=1):
            StepImage.objects.filter(step=step, id=img_id).update(order=index, updated_at=timezone.now())
    return JsonResponse({"ok": True})


@login_required
@require_app_access('process_creator', action='view')
def process_pdf(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    # Optional focused step
    focused_step_id = request.GET.get('focused_step')
    try:
        steps = process.steps.all() if not focused_step_id else process.steps.filter(id=int(focused_step_id))
    except (TypeError, ValueError):
        steps = process.steps.all()
    # Render the print template to HTML
    html_string = render_to_string('process_creator/print.html', {'process': process, 'steps': steps})

    # Helper to resolve static/media URLs for xhtml2pdf
    def link_callback(uri, rel):
        if uri.startswith('/media/'):
            path = os.path.join(settings.MEDIA_ROOT, uri.replace('/media/', ''))
            return path
        if uri.startswith('/static/'):
            # If STATIC_ROOT is set, prefer it; otherwise try finders
            static_root = getattr(settings, 'STATIC_ROOT', None)
            if static_root:
                return os.path.join(static_root, uri.replace('/static/', ''))
        return uri

    # Generate PDF via xhtml2pdf
    from io import BytesIO
    pdf_io = BytesIO()
    pisa.CreatePDF(src=html_string, dest=pdf_io, link_callback=link_callback, encoding='utf-8')
    pdf_io.seek(0)

    # Generate filename with process title and timestamp
    from datetime import datetime
    from urllib.parse import quote
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_title = "".join(c for c in process.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_title = safe_title.replace(' ', '-')
    filename = f"Process-{safe_title}-{timestamp}.pdf"
    
    response = HttpResponse(pdf_io.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response


@login_required
@require_app_access('process_creator', action='view')
def process_word(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    # Optional focused step
    focused_step_id = request.GET.get('focused_step')
    try:
        steps = process.steps.all() if not focused_step_id else process.steps.filter(id=int(focused_step_id))
    except (TypeError, ValueError):
        steps = process.steps.all()
    
    # Get toggle states from request parameters
    # Default to False if not explicitly provided; frontend will pass current toggle states
    show_description = request.GET.get('show_description', 'false').lower() == 'true'
    show_notes = request.GET.get('show_notes', 'false').lower() == 'true'
    show_analysis = request.GET.get('show_analysis', 'false').lower() == 'true'
    show_attachments = request.GET.get('show_attachments', 'false').lower() == 'true'
    show_summary = request.GET.get('show_summary', 'false').lower() == 'true'
    show_pdfs = request.GET.get('show_pdfs', 'false').lower() == 'true'
    
    # Create a new Word document
    doc = Document()
    # Remove default empty first paragraph to avoid accidental blank first page
    try:
        if len(doc.paragraphs) and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
    except Exception:
        pass
    # Normalize section settings to avoid leading blank due to odd-page starts
    try:
        from docx.enum.section import WD_SECTION_START
        for section in doc.sections:
            section.start_type = WD_SECTION_START.NEW_PAGE
            section.different_first_page_header_footer = False
    except Exception:
        pass
    
    # Add title without forcing a title-page behavior
    title = doc.add_paragraph(process.name, style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        title.paragraph_format.page_break_before = False
    except Exception:
        pass
    
    # Add summary (respect toggle)
    if show_summary and process.summary:
        doc.add_heading('Summary', level=1)
        # Use the new Markdown converter for proper formatting
        add_markdown_to_word_doc(doc, process.summary, level=2)
    
    # Add description (if toggle is on)
    if show_description and process.description:
        doc.add_heading('Description', level=1)
        p = doc.add_paragraph(process.description)
        p.paragraph_format.space_after = Inches(0.1)
        p.paragraph_format.line_spacing = 1.15
    
    # Add PDFs before steps if toggle is on
    if show_pdfs:
        pdf_files = []
        for step in steps:
            for file in step.files.all():
                if file.file.name.lower().endswith('.pdf'):
                    pdf_files.append(file)
        
        if pdf_files:
            doc.add_heading('PDF Documents', level=1)
            for pdf_file in pdf_files:
                try:
                    pdf_path = os.path.join(settings.MEDIA_ROOT, str(pdf_file.file))
                    if os.path.exists(pdf_path):
                        pages = _render_pdf_pages_to_images(pdf_path)
                        if pages:
                            for img_io in pages:
                                p = doc.add_paragraph()
                                r = p.add_run()
                                r.add_picture(img_io, width=Inches(6))
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # Fallback: show filename if rendering unavailable
                            doc.add_paragraph(f"PDF: {os.path.basename(pdf_file.file.name)}")
                except Exception:
                    # If anything fails, add a reference
                    doc.add_paragraph(f"PDF: {os.path.basename(pdf_file.file.name)}")
    
    # Add steps (always show steps as they are core content)
    if steps.exists():
        doc.add_heading('Steps', level=1)
        for step in steps:
            # Add step title
            step_heading = doc.add_heading(f'{step.order}. {step.title}', level=2)
            
            # Add step details, and for bullets place images with caption under each image
            if step.details:
                lines = step.details.split('\n')
                bullet_idx = -1
                for raw in lines:
                    line = raw.rstrip('\r')
                    is_bullet = bool(re.match(r'^\s*-\s+', line))
                    if is_bullet:
                        bullet_idx += 1
                    bullet_text_match = re.match(r'^\s*-\s+(.*)$', line)
                    bullet_text = bullet_text_match.group(1) if bullet_text_match else line
                    # If bullet and we should include attachments, render images with caption (caption under image)
                    if is_bullet and show_attachments and step.images.exists():
                        related = [img for img in step.images.all() if img.substep_index == bullet_idx]
                        if related:
                            for img in related:
                                try:
                                    img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                                    if os.path.exists(img_path):
                                        table = doc.add_table(rows=1, cols=1)
                                        table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        cell = table.cell(0, 0)
                                        cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        # Image paragraph
                                        paragraph = cell.paragraphs[0]
                                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        paragraph.paragraph_format.keep_with_next = True
                                        run = paragraph.add_run()
                                        run.add_picture(img_path, width=Inches(6))
                                        # Caption paragraph (bullet text beneath image)
                                        cap = cell.add_paragraph(bullet_text)
                                        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                        cap.paragraph_format.space_before = Pt(6)
                                        cap.paragraph_format.line_spacing = 1.15
                                        cap.paragraph_format.keep_together = True
                                        cap.paragraph_format.keep_with_next = False
                                        # Prevent table row from splitting across pages
                                        from docx.oxml.shared import OxmlElement, qn
                                        tr = table.rows[0]._tr
                                        trPr = tr.get_or_add_trPr()
                                        cantSplit = OxmlElement('w:cantSplit')
                                        trPr.append(cantSplit)
                                        # Slightly larger font for visibility
                                        for r in cap.runs:
                                            r.font.size = Pt(12)
                                        # Add thin border around the image container cell to separate visually
                                        tc = cell._tc
                                        tcPr = tc.get_or_add_tcPr()
                                        tcBorders = OxmlElement('w:tcBorders')
                                        for border_name in ['top', 'left', 'bottom', 'right']:
                                            border = OxmlElement(f'w:{border_name}')
                                            border.set(qn('w:val'), 'single')
                                            border.set(qn('w:sz'), '8')
                                            border.set(qn('w:space'), '0')
                                            border.set(qn('w:color'), 'CCCCCC')
                                            tcBorders.append(border)
                                        tcPr.append(tcBorders)
                                except Exception:
                                    pass
                            # We have rendered bullets as captions under each image; skip separate bullet paragraph
                            continue
                    # Default: render the line as normal paragraph
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Inches(0.05)
                    p.paragraph_format.line_spacing = 1.15
            
            # Any remaining images without substep_index: render after details (respect attachments toggle)
            if show_attachments and step.images.exists():
                for img in step.images.all():
                    if img.substep_index is None:
                        try:
                            img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                            if os.path.exists(img_path):
                                table = doc.add_table(rows=1, cols=1)
                                table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell = table.cell(0, 0)
                                cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                                paragraph = cell.paragraphs[0]
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                run = paragraph.add_run()
                                run.add_picture(img_path, width=Inches(6))
                                from docx.oxml.shared import OxmlElement, qn
                                tc = cell._tc
                                tcPr = tc.get_or_add_tcPr()
                                tcBorders = OxmlElement('w:tcBorders')
                                for border_name in ['top', 'left', 'bottom', 'right']:
                                    border = OxmlElement(f'w:{border_name}')
                                    border.set(qn('w:val'), 'single')
                                    border.set(qn('w:sz'), '12')
                                    border.set(qn('w:space'), '0')
                                    border.set(qn('w:color'), '333333')
                                    tcBorders.append(border)
                                tcPr.append(tcBorders)
                        except Exception:
                            pass
    
    # Add notes (if toggle is on)
    if show_notes and process.notes:
        doc.add_heading('Notes', level=1)
        p = doc.add_paragraph(process.notes)
        p.paragraph_format.space_after = Inches(0.1)
        p.paragraph_format.line_spacing = 1.15
    
    # Add analysis (if toggle is on)
    if show_analysis and process.analysis:
        doc.add_heading('Process Analysis', level=1)
        # Use the new Markdown converter for proper formatting
        add_markdown_to_word_doc(doc, process.analysis, level=2)
    
    # Add Process History section with timestamps
    doc.add_heading('Process History', level=1)
    
    # Process timestamps
    from django.utils import timezone
    process_created = process.created_at.strftime('%B %d, %Y at %I:%M %p')
    process_updated = process.updated_at.strftime('%B %d, %Y at %I:%M %p')
    
    history_p = doc.add_paragraph()
    history_p.add_run('Process: ').bold = True
    history_p.add_run(f'Created {process_created}, Last Updated {process_updated}')
    
    # Step timestamps
    if steps.exists():
        doc.add_heading('Step History', level=2)
        for step in steps:
            step_created = step.created_at.strftime('%B %d, %Y at %I:%M %p')
            step_updated = step.updated_at.strftime('%B %d, %Y at %I:%M %p')
            
            step_p = doc.add_paragraph()
            step_p.add_run(f'Step {step.order}: {step.title} - ').bold = True
            step_p.add_run(f'Created {step_created}, Updated {step_updated}')
            
            # Substep timestamps (for images and files)
            if show_attachments:
                # Image timestamps
                if step.images.exists():
                    doc.add_heading(f'Images for Step {step.order}', level=3)
                    for img in step.images.all():
                        img_uploaded = img.uploaded_at.strftime('%B %d, %Y at %I:%M %p')
                        img_updated = img.updated_at.strftime('%B %d, %Y at %I:%M %p')
                        img_p = doc.add_paragraph()
                        img_p.add_run(f'Image {img.order}: ').bold = True
                        img_p.add_run(f'{os.path.basename(img.image.name)} - Uploaded {img_uploaded}, Updated {img_updated}')
                        if img.substep_index is not None:
                            img_p.add_run(f' (Associated with substep {img.substep_index + 1})')
                
                # File timestamps
                if step.files.exists():
                    doc.add_heading(f'Files for Step {step.order}', level=3)
                    for file in step.files.all():
                        file_uploaded = file.uploaded_at.strftime('%B %d, %Y at %I:%M %p')
                        file_updated = file.updated_at.strftime('%B %d, %Y at %I:%M %p')
                        file_p = doc.add_paragraph()
                        file_p.add_run(f'File {file.order}: ').bold = True
                        file_p.add_run(f'{os.path.basename(file.file.name)} - Uploaded {file_uploaded}, Updated {file_updated}')
                
                # Link timestamps
                if step.links.exists():
                    doc.add_heading(f'Links for Step {step.order}', level=3)
                    for link in step.links.all():
                        link_created = link.created_at.strftime('%B %d, %Y at %I:%M %p')
                        link_updated = link.updated_at.strftime('%B %d, %Y at %I:%M %p')
                        link_p = doc.add_paragraph()
                        link_p.add_run(f'Link {link.order}: ').bold = True
                        link_p.add_run(f'{link.title} - Created {link_created}, Updated {link_updated}')
    
    # Save to BytesIO
    from io import BytesIO
    doc_io = BytesIO()
    # Ensure doc starts content immediately (avoid stray leading section break)
    doc.settings.odd_and_even_pages_header_footer = False
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Generate filename with module (if any), process title and timestamp
    from datetime import datetime
    from urllib.parse import quote
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_process = "".join(c for c in process.name if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '-')
    if process.module:
        safe_module = "".join(c for c in process.module.name if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '-')
        filename = f"{safe_module}-{safe_process}-{timestamp}.docx"
    else:
        filename = f"Process-{safe_process}-{timestamp}.docx"
    
    response = HttpResponse(doc_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response


@login_required
@require_app_access('process_creator', action='view')
def process_stats(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Calculate statistics
    step_count = process.steps.count()
    
    # Count substeps (lines starting with -)
    substep_count = 0
    for step in process.steps.all():
        if step.details:
            # Count lines that start with - (substeps)
            substep_count += len([line for line in step.details.split('\n') if line.strip().startswith('-')])
    
    total_steps = step_count + substep_count
    image_count = sum(step.images.count() for step in process.steps.all())
    
    # Calculate text lengths
    description_length = len(process.description) if process.description else 0
    notes_length = len(process.notes) if process.notes else 0
    summary_length = len(process.summary) if process.summary else 0
    analysis_length = len(process.analysis) if process.analysis else 0
    
    # Calculate total text length from all step details
    step_details_length = 0
    for step in process.steps.all():
        if step.details:
            step_details_length += len(step.details)
    
    # Total text length (all text content combined)
    total_text_length = description_length + notes_length + summary_length + analysis_length + step_details_length
    
    # Format dates
    from django.utils import timezone
    created_at = process.created_at.strftime('%B %d, %Y at %I:%M %p')
    updated_at = process.updated_at.strftime('%B %d, %Y at %I:%M %p')
    
    stats = {
        'created_at': created_at,
        'updated_at': updated_at,
        'step_count': step_count,
        'substep_count': substep_count,
        'total_steps': total_steps,
        'image_count': image_count,
        'description_length': description_length,
        'notes_length': notes_length,
        'summary_length': summary_length,
        'analysis_length': analysis_length,
        'step_details_length': step_details_length,
        'total_text_length': total_text_length,
    }
    
    return JsonResponse(stats)


def call_openai_api(prompt, model="gpt-4o-mini", max_tokens=2000):
    """Helper function to call OpenAI API and track usage"""
    try:
        if not settings.OPENAI_API_KEY:
            return {
                'success': False,
                'error': 'OpenAI API key not configured. Please set OPENAI_API_KEY environment variable.',
                'tokens_used': 0,
                'cost': Decimal('0.00')
            }
        
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.7
        )
        
        return {
            'success': True,
            'content': response.choices[0].message.content,
            'tokens_used': response.usage.total_tokens,
            'cost': calculate_cost(response.usage.total_tokens, model)
        }
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'tokens_used': 0,
            'cost': Decimal('0.00')
        }


def calculate_cost(tokens, model):
    """Calculate cost based on token usage and model"""
    # Pricing per 1K tokens (as of 2024)
    pricing = {
        'gpt-4o-mini': {'input': 0.00015, 'output': 0.0006},
        'gpt-4o': {'input': 0.005, 'output': 0.015},
        'gpt-3.5-turbo': {'input': 0.0015, 'output': 0.002}
    }
    
    if model not in pricing:
        model = 'gpt-4o-mini'
    
    # Rough estimate - assume 50/50 input/output split
    cost_per_token = (pricing[model]['input'] + pricing[model]['output']) / 2
    return Decimal(str(tokens * cost_per_token / 1000)).quantize(Decimal('0.000001'))


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def ai_generate_summary(request, pk: int):
    """Generate AI summary for a process"""
    process = get_object_or_404(Process, pk=pk)
    
    # Debug: Check if API key is available
    if not settings.OPENAI_API_KEY:
        return JsonResponse({
            'success': False, 
            'error': 'OpenAI API key not configured. Please set OPENAI_API_KEY environment variable.'
        })
    
    # Get custom instructions or use default
    data = json.loads(request.body)
    instructions = data.get('instructions', process.summary_instructions)
    
    # Build prompt with process data
    prompt = f"""{instructions}

Process Details:
Name: {process.name}
Description: {process.description or 'No description provided'}

Steps:
"""
    
    for step in process.steps.all():
        prompt += f"{step.order}. {step.title}\n"
        if step.details:
            prompt += f"   Details: {step.details}\n"
        if step.images.exists():
            prompt += f"   Images: {step.images.count()} screenshot(s)\n"
        prompt += "\n"
    
    if process.notes:
        prompt += f"Notes: {process.notes}\n"
    
    # Call OpenAI API
    result = call_openai_api(prompt, max_tokens=500)
    
    if result['success']:
        # Update process
        process.summary = result['content']
        process.summary_instructions = instructions
        process.last_ai_update = timezone.now()
        process.save()
        
        # Log interaction
        AIInteraction.objects.create(
            process=process,
            interaction_type='summary',
            prompt_sent=prompt,
            response_received=result['content'],
            tokens_used=result['tokens_used'],
            cost=result['cost']
        )
        
        return JsonResponse({
            'success': True,
            'summary': result['content']
        })
    else:
        return JsonResponse({
            'success': False,
            'error': result['error']
        })


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def ai_analyze_process(request, pk: int):
    """Generate AI analysis for process improvement"""
    process = get_object_or_404(Process, pk=pk)
    
    # Debug: Check if API key is available
    if not settings.OPENAI_API_KEY:
        return JsonResponse({
            'success': False, 
            'error': 'OpenAI API key not configured. Please set OPENAI_API_KEY environment variable.'
        })
    
    # Get custom instructions or use default
    data = json.loads(request.body)
    instructions = data.get('instructions', process.analysis_instructions)
    
    # Build comprehensive prompt
    prompt = f"""You are a business process improvement consultant. {instructions}

Process Details:
Name: {process.name}
Description: {process.description or 'No description provided'}

Current Process Steps:
"""
    
    for step in process.steps.all():
        prompt += f"{step.order}. {step.title}\n"
        if step.details:
            prompt += f"   Details: {step.details}\n"
        if step.images.exists():
            prompt += f"   Images: {step.images.count()} screenshot(s)\n"
        prompt += "\n"
    
    if process.notes:
        prompt += f"Additional Notes: {process.notes}\n"
    
    prompt += f"\nCurrent Summary: {process.summary or 'No summary available'}\n"
    
    prompt += """
Please provide a detailed analysis following the instructions above. Structure your response with clear headings and actionable recommendations.
"""
    
    # Call OpenAI API
    result = call_openai_api(prompt, max_tokens=2000)
    
    if result['success']:
        # Update process
        process.analysis = result['content']
        process.analysis_instructions = instructions
        process.last_ai_update = timezone.now()
        process.save()
        
        # Log interaction
        AIInteraction.objects.create(
            process=process,
            interaction_type='analysis',
            prompt_sent=prompt,
            response_received=result['content'],
            tokens_used=result['tokens_used'],
            cost=result['cost']
        )
        
        return JsonResponse({
            'success': True,
            'analysis': result['content']
        })
    else:
        return JsonResponse({
            'success': False,
            'error': result['error']
        })

# Bulk Operations
@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def bulk_summary(request):
    """Generate summary for multiple processes"""
    try:
        data = json.loads(request.body)
        process_ids = data.get('process_ids', [])
        
        if not process_ids:
            return JsonResponse({'success': False, 'error': 'No processes selected'})
        
        # Get all selected processes
        processes = Process.objects.filter(id__in=process_ids).order_by('order')
        if not processes.exists():
            return JsonResponse({'success': False, 'error': 'No valid processes found'})
        
        # Combine all process data
        combined_data = []
        for process in processes:
            process_info = {
                'name': process.name,
                'description': process.description,
                'steps': []
            }
            
            for step in process.steps.all().order_by('order'):
                step_info = {
                    'title': step.title,
                    'details': step.details
                }
                process_info['steps'].append(step_info)
            
            combined_data.append(process_info)
        
        # Use the default summary instructions
        summary_instructions = "You are given a list of steps, bullet points, or fragmented notes. Your task is to transform them into a single professional, coherent paragraph. Do not repeat the steps as a list. Instead, weave them into smooth, natural prose that reads as if written by a skilled professional writer. Maintain accuracy, logical flow, and clarity. The output must always be a polished paragraph summary, never a bullet list."
        
        # Create comprehensive prompt
        prompt = f"{summary_instructions}\n\nProcess Data:\n{json.dumps(combined_data, indent=2)}"
        
        # Call OpenAI API
        result = call_openai_api(prompt, model='gpt-4o-mini')
        
        if result['success']:
            return JsonResponse({
                'success': True,
                'summary': result['content']
            })
        else:
            return JsonResponse({
                'success': False,
                'error': result['error']
            })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Error generating bulk summary: {str(e)}'
        })


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def bulk_analyze(request):
    """Generate analysis for multiple processes"""
    try:
        data = json.loads(request.body)
        process_ids = data.get('process_ids', [])
        
        if not process_ids:
            return JsonResponse({'success': False, 'error': 'No processes selected'})
        
        # Get all selected processes
        processes = Process.objects.filter(id__in=process_ids).order_by('order')
        if not processes.exists():
            return JsonResponse({'success': False, 'error': 'No valid processes found'})
        
        # Combine all process data
        combined_data = []
        for process in processes:
            process_info = {
                'name': process.name,
                'description': process.description,
                'notes': process.notes,
                'steps': []
            }
            
            for step in process.steps.all().order_by('order'):
                step_info = {
                    'title': step.title,
                    'details': step.details
                }
                process_info['steps'].append(step_info)
            
            combined_data.append(process_info)
        
        # Use the default analysis instructions
        analysis_instructions = """You are a senior operations analyst. You will receive a business process as input (steps, notes, and context). Produce a professional report titled "Training Analysis" that managers and frontline staff can act on.
Follow these rules:
- Write clearly and concisely; avoid jargon. Use confident, direct language.
- Do NOT restate the steps as bullets. Summarize the process in flowing prose.
- Use Markdown headings exactly as specified below.
- Give specific, actionable recommendations with clear benefits and effort levels.
- Where details are missing, make reasonable assumptions and state them briefly.
=== INPUT START ===
{PASTE PROCESS STEPS / NOTES / CONTEXT HERE}
=== INPUT END ===
=== OUTPUT FORMAT (Markdown) ===
# Training Analysis
## Executive Summary
A 4–6 sentence overview of the process, the primary objective, the current state, and the most important recommended changes and expected benefits.
## Process Narrative (Plain-Language Summary)
6–10 sentences that narrate how the process works from start to finish. Use paragraph form only (no bullets). Emphasize flow, handoffs, tools used, approvals, and timing.
## Assumptions
- 2–5 short bullets listing key assumptions you made due to missing information.
## Strengths
- 4–8 bullets. Each bullet: what works well + why it matters (impact on quality, speed, cost, compliance, safety, or CX).
## Weaknesses
- 4–8 bullets. Each bullet: the issue + consequence (e.g., rework, delays, risk, cost).
## Improvement Opportunities — Current System (No new software)
Provide 5–10 specific, low-friction changes using existing tools, roles, and policies. For each item, use this format in one bullet:
- **[Title]** — What to change (1–2 sentences). **Benefit:** expected outcome with a measurable indicator (e.g., "reduce cycle time by ~15%"). **Effort:** Low/Med/High. **Owner:** role. **Risk/Mitigation:** brief.
## Immediate Application-Level Enhancements (Quick to implement in an app/workflow)
Provide 3–7 changes that can be implemented quickly via simple configuration, scripts, forms, validations, templates, or notifications. For each item, use the same format:
- **[Title]** — What to implement (1–2 sentences). **Benefit:** measurable. **Effort:** Low/Med/High. **Owner:** role. **Risk/Mitigation:** brief.
## Implementation Roadmap
Organize recommendations into phases with rationale:
- **Now (0–30 days):** 3–6 highest-ROI, low-effort actions.
- **Next (30–90 days):** 3–6 medium-effort actions that compound earlier gains.
- **Later (90+ days):** 2–4 higher-effort changes that deliver strategic value.
## Metrics & Monitoring
List 4–8 KPIs with target direction and cadence. For each KPI: **Name**, **Target/Direction**, **Data Source**, **Review Cadence** (e.g., weekly), **Owner**.
## Risks & Mitigations
3–6 material risks across people/process/tech/compliance and how to mitigate each.
## Final Recommendation
A 4–6 sentence closing paragraph summarizing the case for change, expected benefits, and the immediate next steps.
=== STYLE GUARDRAILS ===
- Professional tone. Tight sentences. No filler, no hype.
- Use paragraph form for the narrative sections; bullets only where specified.
- Quantify benefits or ranges when plausible. Avoid vague claims.
- Do not include meta-commentary or instructions in the output."""
        
        # Create comprehensive prompt
        prompt = f"{analysis_instructions}\n\nProcess Data:\n{json.dumps(combined_data, indent=2)}"
        
        # Call OpenAI API
        result = call_openai_api(prompt, model='gpt-4o-mini')
        
        if result['success']:
            return JsonResponse({
                'success': True,
                'analysis': result['content']
            })
        else:
            return JsonResponse({
                'success': False,
                'error': result['error']
            })
            
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': f'Error generating bulk analysis: {str(e)}'
        })


@login_required
@require_app_access('process_creator', action='view')
def bulk_pdf(request):
    """Generate PDF for multiple processes"""
    process_ids = request.GET.getlist('ids')
    if not process_ids:
        return HttpResponse('No processes selected', status=400)
    
    # Filter by module if specified
    selected_module_id = request.GET.get('module')
    if selected_module_id:
        try:
            selected_module = Module.objects.get(id=selected_module_id)
            processes = Process.objects.filter(id__in=process_ids, module=selected_module).order_by('order')
        except Module.DoesNotExist:
            processes = Process.objects.filter(id__in=process_ids).order_by('order')
    else:
        processes = Process.objects.filter(id__in=process_ids).order_by('order')
    
    if not processes.exists():
        return HttpResponse('No valid processes found', status=400)
    
    # Get history data if included
    history_data = []
    if request.GET.get('include_history') == 'true':
        try:
            history_json = request.GET.get('history_data', '[]')
            history_data = json.loads(history_json)
        except (json.JSONDecodeError, TypeError):
            history_data = []
    
    # Render template
    html_string = render_to_string('process_creator/bulk_print.html', {
        'processes': processes,
        'history_data': history_data
    })
    
    def link_callback(uri, rel):
        sUrl = settings.STATIC_URL
        sRoot = settings.STATIC_ROOT
        mUrl = settings.MEDIA_URL
        mRoot = settings.MEDIA_ROOT

        if uri.startswith(mUrl):
            path = os.path.join(mRoot, uri.replace(mUrl, ""))
        elif uri.startswith(sUrl):
            path = os.path.join(sRoot, uri.replace(sUrl, ""))
        else:
            path = os.path.join(settings.BASE_DIR, uri)

        if not os.path.isfile(path):
            return uri
        return path

    response = HttpResponse(content_type='application/pdf')
    pdf_io = BytesIO()
    pisa.CreatePDF(src=html_string, dest=pdf_io, link_callback=link_callback, encoding='utf-8')
    pdf_io.seek(0)

    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    filename = f"Bulk-Process-Report-{timestamp}.pdf"
    
    response = HttpResponse(pdf_io.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response


@login_required
@require_app_access('process_creator', action='view')
def bulk_word(request):
    """Generate Word document for multiple processes"""
    process_ids = request.GET.getlist('ids')
    if not process_ids:
        return HttpResponse('No processes selected', status=400)
    
    # Filter by module if specified
    selected_module_id = request.GET.get('module')
    if selected_module_id:
        try:
            selected_module = Module.objects.get(id=selected_module_id)
            processes = Process.objects.filter(id__in=process_ids, module=selected_module).order_by('order')
        except Module.DoesNotExist:
            processes = Process.objects.filter(id__in=process_ids).order_by('order')
    else:
        processes = Process.objects.filter(id__in=process_ids).order_by('order')
    
    if not processes.exists():
        return HttpResponse('No valid processes found', status=400)
    
    # Get history data if included
    history_data = []
    if request.GET.get('include_history') == 'true':
        try:
            history_json = request.GET.get('history_data', '[]')
            history_data = json.loads(history_json)
        except (json.JSONDecodeError, TypeError):
            history_data = []
    
    doc = Document()
    # Remove default empty first paragraph to avoid accidental blank first page
    try:
        if len(doc.paragraphs) and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
    except Exception:
        pass
    # Normalize section settings to avoid leading blank due to odd-page starts
    try:
        from docx.enum.section import WD_SECTION_START
        for section in doc.sections:
            section.start_type = WD_SECTION_START.NEW_PAGE
            section.different_first_page_header_footer = False
    except Exception:
        pass
    
    # Add title without using the Title style (which can create a title page)
    title = doc.add_paragraph('Bulk Process Report', style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        title.paragraph_format.page_break_before = False
    except Exception:
        pass
    
    # Read toggle states (module-level controls)
    show_description = request.GET.get('show_description', 'false').lower() == 'true'
    show_notes = request.GET.get('show_notes', 'false').lower() == 'true'
    show_analysis = request.GET.get('show_analysis', 'false').lower() == 'true'
    show_summary = request.GET.get('show_summary', 'false').lower() == 'true'
    show_attachments = request.GET.get('show_attachments', 'false').lower() == 'true'
    show_pdfs = request.GET.get('show_pdfs', 'false').lower() == 'true'

    # Add each process
    for i, process in enumerate(processes, 1):
        if i > 1:
            doc.add_page_break()
        
        # Process name
        process_heading = doc.add_heading(f'{i}. {process.name}', level=1)
        
        # Add PDFs before steps if toggle is on
        if show_pdfs:
            pdf_files = []
            for step in process.steps.all():
                for file in step.files.all():
                    if file.file.name.lower().endswith('.pdf'):
                        pdf_files.append(file)
            
            if pdf_files:
                doc.add_heading('PDF Documents', level=2)
                for pdf_file in pdf_files:
                    try:
                        pdf_path = os.path.join(settings.MEDIA_ROOT, str(pdf_file.file))
                        if os.path.exists(pdf_path):
                            pages = _render_pdf_pages_to_images(pdf_path)
                            if pages:
                                for img_io in pages:
                                    p = doc.add_paragraph()
                                    r = p.add_run()
                                    r.add_picture(img_io, width=Inches(6))
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                doc.add_paragraph(f"PDF: {os.path.basename(pdf_file.file.name)}")
                    except Exception:
                        # If PDF can't be embedded, add a reference
                        doc.add_paragraph(f"PDF: {os.path.basename(pdf_file.file.name)}")
        
        # Summary (Markdown -> Word)
        if show_summary and process.summary:
            doc.add_heading('Summary', level=2)
            add_markdown_to_word_doc(doc, process.summary, level=3)
        
        # Description
        if show_description and process.description:
            doc.add_heading('Description', level=2)
            p = doc.add_paragraph(process.description)
            p.paragraph_format.space_after = Inches(0.1)
            p.paragraph_format.line_spacing = 1.15
        
        # Steps
        if process.steps.exists():
            doc.add_heading('Steps', level=2)
            for step in process.steps.all():
                step_heading = doc.add_heading(f'{step.order}. {step.title}', level=3)
                
                if step.details:
                    lines = step.details.split('\n')
                    bullet_idx = -1
                    for raw in lines:
                        line = raw.rstrip('\r')
                        is_bullet = bool(re.match(r'^\s*-\s+', line))
                        if is_bullet:
                            bullet_idx += 1
                        bullet_text_match = re.match(r'^\s*-\s+(.*)$', line)
                        bullet_text = bullet_text_match.group(1) if bullet_text_match else line
                        if is_bullet and show_attachments and step.images.exists():
                            related = [img for img in step.images.all() if img.substep_index == bullet_idx]
                            if related:
                                for img in related:
                                    try:
                                        img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                                        if os.path.exists(img_path):
                                            table = doc.add_table(rows=1, cols=1)
                                            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            cell = table.cell(0, 0)
                                            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            paragraph = cell.paragraphs[0]
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            paragraph.paragraph_format.keep_with_next = True
                                            run = paragraph.add_run()
                                            run.add_picture(img_path, width=Inches(6))
                                            cap = cell.add_paragraph(bullet_text)
                                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            cap.paragraph_format.space_before = Pt(6)
                                            cap.paragraph_format.line_spacing = 1.15
                                            cap.paragraph_format.keep_together = True
                                            cap.paragraph_format.keep_with_next = False
                                            # Prevent table row from splitting across pages
                                            from docx.oxml.shared import OxmlElement, qn
                                            tr = table.rows[0]._tr
                                            trPr = tr.get_or_add_trPr()
                                            cantSplit = OxmlElement('w:cantSplit')
                                            trPr.append(cantSplit)
                                            for r in cap.runs:
                                                r.font.size = Pt(12)
                                            tc = cell._tc
                                            tcPr = tc.get_or_add_tcPr()
                                            tcBorders = OxmlElement('w:tcBorders')
                                            for border_name in ['top', 'left', 'bottom', 'right']:
                                                border = OxmlElement(f'w:{border_name}')
                                                border.set(qn('w:val'), 'single')
                                                border.set(qn('w:sz'), '8')
                                                border.set(qn('w:space'), '0')
                                                border.set(qn('w:color'), 'CCCCCC')
                                                tcBorders.append(border)
                                            tcPr.append(tcBorders)
                                    except Exception:
                                        pass
                                continue
                        p = doc.add_paragraph(line)
                        p.paragraph_format.space_after = Inches(0.05)
                        p.paragraph_format.line_spacing = 1.15
                # Any remaining images without substep_index
                if show_attachments and step.images.exists():
                    for img in step.images.all():
                        if img.substep_index is None:
                            try:
                                img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                                if os.path.exists(img_path):
                                    table = doc.add_table(rows=1, cols=1)
                                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cell = table.cell(0, 0)
                                    cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture(img_path, width=Inches(6))
                                    from docx.oxml.shared import OxmlElement, qn
                                    tc = cell._tc
                                    tcPr = tc.get_or_add_tcPr()
                                    tcBorders = OxmlElement('w:tcBorders')
                                    for border_name in ['top', 'left', 'bottom', 'right']:
                                        border = OxmlElement(f'w:{border_name}')
                                        border.set(qn('w:val'), 'single')
                                        border.set(qn('w:sz'), '12')
                                        border.set(qn('w:space'), '0')
                                        border.set(qn('w:color'), '333333')
                                        tcBorders.append(border)
                                    tcPr.append(tcBorders)
                            except Exception:
                                pass
        
        # Notes
        if show_notes and process.notes:
            doc.add_heading('Notes', level=2)
            p = doc.add_paragraph(process.notes)
            p.paragraph_format.space_after = Inches(0.1)
            p.paragraph_format.line_spacing = 1.15
        
        # Analysis (Markdown -> Word)
        if show_analysis and process.analysis:
            doc.add_heading('Process Analysis', level=2)
            add_markdown_to_word_doc(doc, process.analysis, level=3)
        
        # Add Process History section with timestamps for this process
        doc.add_heading('Process History', level=2)
        
        # Process timestamps
        from django.utils import timezone
        process_created = process.created_at.strftime('%B %d, %Y at %I:%M %p')
        process_updated = process.updated_at.strftime('%B %d, %Y at %I:%M %p')
        
        history_p = doc.add_paragraph()
        history_p.add_run('Process: ').bold = True
        history_p.add_run(f'Created {process_created}, Last Updated {process_updated}')
        
        # Step timestamps
        if process.steps.exists():
            doc.add_heading('Step History', level=3)
            for step in process.steps.all():
                step_created = step.created_at.strftime('%B %d, %Y at %I:%M %p')
                step_updated = step.updated_at.strftime('%B %d, %Y at %I:%M %p')
                
                step_p = doc.add_paragraph()
                step_p.add_run(f'Step {step.order}: {step.title} - ').bold = True
                step_p.add_run(f'Created {step_created}, Updated {step_updated}')
                
                # Substep timestamps (for images and files)
                if show_attachments:
                    # Image timestamps
                    if step.images.exists():
                        doc.add_heading(f'Images for Step {step.order}', level=4)
                        for img in step.images.all():
                            img_uploaded = img.uploaded_at.strftime('%B %d, %Y at %I:%M %p')
                            img_updated = img.updated_at.strftime('%B %d, %Y at %I:%M %p')
                            img_p = doc.add_paragraph()
                            img_p.add_run(f'Image {img.order}: ').bold = True
                            img_p.add_run(f'{os.path.basename(img.image.name)} - Uploaded {img_uploaded}, Updated {img_updated}')
                            if img.substep_index is not None:
                                img_p.add_run(f' (Associated with substep {img.substep_index + 1})')
                    
                    # File timestamps
                    if step.files.exists():
                        doc.add_heading(f'Files for Step {step.order}', level=4)
                        for file in step.files.all():
                            file_uploaded = file.uploaded_at.strftime('%B %d, %Y at %I:%M %p')
                            file_updated = file.updated_at.strftime('%B %d, %Y at %I:%M %p')
                            file_p = doc.add_paragraph()
                            file_p.add_run(f'File {file.order}: ').bold = True
                            file_p.add_run(f'{os.path.basename(file.file.name)} - Uploaded {file_uploaded}, Updated {file_updated}')
                    
                    # Link timestamps
                    if step.links.exists():
                        doc.add_heading(f'Links for Step {step.order}', level=4)
                        for link in step.links.all():
                            link_created = link.created_at.strftime('%B %d, %Y at %I:%M %p')
                            link_updated = link.updated_at.strftime('%B %d, %Y at %I:%M %p')
                            link_p = doc.add_paragraph()
                            link_p.add_run(f'Link {link.order}: ').bold = True
                            link_p.add_run(f'{link.title} - Created {link_created}, Updated {link_updated}')
    
    # Add history section if history data is provided
    if history_data:
        doc.add_page_break()
        doc.add_heading('History', level=1)
        
        for i, history_item in enumerate(history_data, 1):
            # History item title
            history_heading = doc.add_heading(f'{i}. {history_item.get("type", "Unknown")} - {history_item.get("date", "")}', level=2)
            
            # History content (Markdown -> Word)
            content = history_item.get('content', '')
            if content:
                add_markdown_to_word_doc(doc, content, level=3)
    
    doc_io = BytesIO()
    doc.settings.odd_and_even_pages_header_footer = False
    doc.save(doc_io)
    doc_io.seek(0)
    
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    # Prefer module name in filename if all processes share one module or module filter provided
    module_name = None
    selected_module_id = request.GET.get('module')
    if selected_module_id:
        try:
            module_obj = Module.objects.get(id=selected_module_id)
            module_name = module_obj.name
        except Module.DoesNotExist:
            module_name = None
    if module_name is None:
        # Derive common module if all selected processes share same module
        modules = list({p.module.name for p in processes if p.module is not None})
        if len(modules) == 1:
            module_name = modules[0]
    if module_name:
        safe_module = "".join(c for c in module_name if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '-')
        filename = f"{safe_module}-{timestamp}.docx"
    else:
        filename = f"Bulk-Process-Report-{timestamp}.docx"
    
    response = HttpResponse(doc_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response

# Create your views here.

# --- Minimal stubs for Templates & Jobs (will be enhanced) ---
@login_required
@require_app_access('process_creator', action='view')
def template_list(request):
    qs = ProcessTemplate.objects.select_related('module', 'source_process').order_by('-updated_at')
    modules = Module.objects.all()
    selected_module = None
    module_id = request.GET.get('module')
    if module_id:
        try:
            selected_module = Module.objects.get(id=module_id)
            qs = qs.filter(module=selected_module)
        except Module.DoesNotExist:
            pass
    if request.GET.get('all') != '1':
        qs = qs.filter(is_active=True)
    return render(request, 'process_creator/templates_list.html', {
        'templates': qs,
        'showing_all': request.GET.get('all') == '1',
        'JOB_LABEL': JOB_LABEL,
        'modules': modules,
        'selected_module': selected_module,
    })


@login_required
@require_app_access('process_creator', action='view')
def template_detail(request, tpl_id: int):
    template = get_object_or_404(
        ProcessTemplate.objects.select_related('module', 'source_process').prefetch_related('steps'), id=tpl_id
    )
    source_steps = list(template.source_process.steps.all()) if template.source_process_id else []
    source_by_order = {s.order: s for s in source_steps}
    return render(request, 'process_creator/template_detail.html', {
        'template': template,
        'source_by_order': source_by_order,
        'JOB_LABEL': JOB_LABEL,
    })


@login_required
@require_app_access('process_creator', action='view')
def job_list(request):
    jobs = Job.objects.select_related('template', 'template__module').order_by('-created_at')
    return render(request, 'process_creator/job_list.html', {'jobs': jobs, 'JOB_LABEL': JOB_LABEL})


@login_required
@require_app_access('process_creator', action='edit')
def job_create(request, tpl_id: int):
    template = get_object_or_404(ProcessTemplate.objects.prefetch_related('steps'), id=tpl_id, is_active=True)
    if request.method == 'POST':
        name = (request.POST.get('name') or '').strip()
        assigned_to_id = request.POST.get('assigned_to')
        if not name:
            return render(request, 'process_creator/job_create.html', {'template': template, 'error': 'Name is required', 'JOB_LABEL': JOB_LABEL})
        assignee = None
        if assigned_to_id:
            from django.contrib.auth import get_user_model
            User = get_user_model()
            assignee = User.objects.filter(id=assigned_to_id).first()
        with transaction.atomic():
            job = Job.objects.create(template=template, name=name, assigned_to=assignee, template_version_at_create=template.version)
            steps = list(template.steps.all().order_by('order', 'id'))
            JobStep.objects.bulk_create([JobStep(job=job, order=s.order, title=s.title, details=s.details or '') for s in steps])
            # Create bullet-level subtasks
            from .models import JobSubtask
            job_steps = {js.order: js for js in job.steps.all()}
            for s in steps:
                js = job_steps.get(s.order)
                if not js:
                    continue
                if s.details:
                    idx = 0
                    for line in s.details.split('\n'):
                        if line.strip().startswith('- '):
                            JobSubtask.objects.create(job_step=js, order=idx, text=line.strip()[2:].strip())
                            idx += 1
        return redirect('process_creator:job_detail', job_id=job.id)
    else:
        from django.contrib.auth import get_user_model
        User = get_user_model()
        users = User.objects.all().order_by('username')
        return render(request, 'process_creator/job_create.html', {'template': template, 'users': users, 'JOB_LABEL': JOB_LABEL})


@login_required
@require_app_access('process_creator', action='view')
def job_detail(request, job_id: int):
    job = get_object_or_404(Job.objects.select_related('template', 'template__module').prefetch_related('steps', 'steps__subtasks', 'steps__images'), id=job_id)
    source_steps = list(job.template.source_process.steps.prefetch_related('images').all()) if job.template and job.template.source_process_id else []
    source_by_order = {s.order: s for s in source_steps}
    source_image_counts = {s.order: (s.images.all().count()) for s in source_steps}
    completed = sum(1 for s in job.steps.all() if s.status == 'completed')
    total = job.steps.count()
    blocked = sum(1 for s in job.steps.all() if s.status == 'blocked')
    total_subtasks = 0
    done_subtasks = 0
    step_percents = {}
    for s in job.steps.all():
        st_total = s.subtasks.count()
        if st_total:
            st_done = s.subtasks.filter(completed=True).count()
            step_percents[s.id] = int(round((st_done / st_total) * 100))
            total_subtasks += st_total
            done_subtasks += st_done
        else:
            step_percents[s.id] = 100 if s.status == 'completed' else 0
    overall_percent = int(round((done_subtasks / total_subtasks) * 100)) if total_subtasks else int(round((completed / total) * 100)) if total else 0
    template_newer = job.template.version > job.template_version_at_create
    return render(request, 'process_creator/job_detail.html', {
        'job': job,
        'source_by_order': source_by_order,
        'progress': {'completed': completed, 'total': total, 'blocked': blocked},
        'JOB_LABEL': JOB_LABEL,
        'template_newer': template_newer,
        'overall_percent': overall_percent,
        'step_percents': step_percents,
        'source_image_counts': source_image_counts,
    })


@login_required
@require_app_access('process_creator', action='edit')
def job_start(request, job_id: int):
    job = get_object_or_404(Job, id=job_id)
    if job.status == 'not_started':
        job.status = 'in_progress'
        job.started_at = timezone.now()
        job.save(update_fields=['status', 'started_at', 'updated_at'])
    return redirect('process_creator:job_detail', job_id=job.id)


def _update_job_step(job: Job, step: JobStep, status: str):
    now = timezone.now()
    step.status = status
    if status == 'in_progress' and not step.started_at:
        step.started_at = now
    if status == 'completed':
        step.completed_at = now
    step.save(update_fields=['status', 'started_at', 'completed_at', 'updated_at'])
    if status == 'completed':
        if all(s.status == 'completed' for s in job.steps.all()):
            job.status = 'completed'
            job.completed_at = now
            job.save(update_fields=['status', 'completed_at', 'updated_at'])


@login_required
@require_app_access('process_creator', action='edit')
def job_step_start(request, job_id: int, job_step_id: int):
    job = get_object_or_404(Job, id=job_id)
    step = get_object_or_404(JobStep, id=job_step_id, job=job)
    _update_job_step(job, step, 'in_progress')
    return redirect('process_creator:job_detail', job_id=job.id)


@login_required
@require_app_access('process_creator', action='edit')
def job_step_complete(request, job_id: int, job_step_id: int):
    job = get_object_or_404(Job, id=job_id)
    step = get_object_or_404(JobStep, id=job_step_id, job=job)
    _update_job_step(job, step, 'completed')
    return redirect('process_creator:job_detail', job_id=job.id)


@login_required
@require_app_access('process_creator', action='edit')
def job_step_block(request, job_id: int, job_step_id: int):
    job = get_object_or_404(Job, id=job_id)
    step = get_object_or_404(JobStep, id=job_step_id, job=job)
    _update_job_step(job, step, 'blocked')
    return redirect('process_creator:job_detail', job_id=job.id)


@login_required
@require_app_access('process_creator', action='edit')
def job_step_unblock(request, job_id: int, job_step_id: int):
    job = get_object_or_404(Job, id=job_id)
    step = get_object_or_404(JobStep, id=job_step_id, job=job)
    new_status = 'in_progress' if step.started_at else 'pending'
    _update_job_step(job, step, new_status)
    return redirect('process_creator:job_detail', job_id=job.id)


@login_required
@require_app_access('process_creator', action='view')
def job_print(request, job_id: int):
    job = get_object_or_404(Job.objects.select_related('template', 'template__source_process'), id=job_id)
    steps = job.steps.all().order_by('order', 'id')
    return render(request, 'process_creator/job_print.html', {'job': job, 'steps': steps, 'JOB_LABEL': JOB_LABEL})


@login_required
@require_app_access('process_creator', action='view')
def job_pdf(request, job_id: int):
    job = get_object_or_404(Job.objects.select_related('template', 'template__source_process'), id=job_id)
    steps = job.steps.all().order_by('order', 'id')
    html_string = render_to_string('process_creator/job_print.html', {'job': job, 'steps': steps, 'JOB_LABEL': JOB_LABEL})
    def link_callback(uri, rel):
        if uri.startswith('/media/'):
            path = os.path.join(settings.MEDIA_ROOT, uri.replace('/media/', ''))
            return path
        if uri.startswith('/static/'):
            static_root = getattr(settings, 'STATIC_ROOT', None)
            if static_root:
                return os.path.join(static_root, uri.replace('/static/', ''))
        return uri
    pdf_io = BytesIO()
    pisa.CreatePDF(src=html_string, dest=pdf_io, link_callback=link_callback, encoding='utf-8')
    pdf_io.seek(0)
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')
    safe_title = ''.join(c for c in job.name if c.isalnum() or c in (' ', '-', '_')).rstrip().replace(' ', '-')
    filename = f"{JOB_LABEL}-{safe_title}-{timestamp}.pdf"
    response = HttpResponse(pdf_io.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response


@login_required
@require_app_access('process_creator', action='view')
def job_word(request, job_id: int):
    job = get_object_or_404(Job.objects.select_related('template', 'template__source_process'), id=job_id)
    steps = job.steps.all().order_by('order', 'id')
    show_attachments = request.GET.get('show_attachments', 'false').lower() == 'true'
    doc = Document()
    try:
        if len(doc.paragraphs) and not doc.paragraphs[0].text.strip():
            p = doc.paragraphs[0]
            p._element.getparent().remove(p._element)
    except Exception:
        pass
    try:
        from docx.enum.section import WD_SECTION_START
        for section in doc.sections:
            section.start_type = WD_SECTION_START.NEW_PAGE
            section.different_first_page_header_footer = False
    except Exception:
        pass
    title = doc.add_paragraph(f"{JOB_LABEL}: {job.name}", style='Heading 1')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        title.paragraph_format.page_break_before = False
    except Exception:
        pass
    # Overall percent based on subtasks
    total_subtasks = 0
    done_subtasks = 0
    for s in steps:
        total_subtasks += s.subtasks.count()
        done_subtasks += s.subtasks.filter(completed=True).count()
    overall_percent = int(round((done_subtasks / total_subtasks) * 100)) if total_subtasks else int(round((sum(1 for s in steps if s.status=='completed') / (steps.count() or 1)) * 100))

    # Status summary
    doc.add_paragraph(f"Status: {job.get_status_display()} — Completion: {overall_percent}%")

    if steps.exists():
        doc.add_heading('Steps', level=1)
        source_steps = list(job.template.source_process.steps.all()) if job.template and job.template.source_process_id else []
        source_by_order = {s.order: s for s in source_steps}
        for step in steps:
            st_total = step.subtasks.count()
            st_done = step.subtasks.filter(completed=True).count()
            st_pct = int(round((st_done / st_total) * 100)) if st_total else (100 if step.status == 'completed' else 0)
            doc.add_heading(f'{step.order}. {step.title} — {st_pct}%', level=2)
            if step.details:
                lines = step.details.split('\n')
                bullet_idx = -1
                for raw in lines:
                    line = raw.rstrip('\r')
                    is_bullet = bool(re.match(r'^\s*-\s+', line))
                    if is_bullet:
                        bullet_idx += 1
                    bullet_text_match = re.match(r'^\s*-\s+(.*)$', line)
                    bullet_text = bullet_text_match.group(1) if bullet_text_match else line
                    if is_bullet and show_attachments:
                        src = source_by_order.get(step.order)
                        if src and src.images.exists():
                            related = [img for img in src.images.all() if img.substep_index == bullet_idx]
                            if related:
                                for img in related:
                                    try:
                                        img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                                        if os.path.exists(img_path):
                                            # Get image dimensions to scale appropriately
                                            from PIL import Image
                                            with Image.open(img_path) as img_pil:
                                                width_px, height_px = img_pil.size
                                                # Calculate width based on image size relative to typical screen (1920px)
                                                # Full screen = 6 inches, smaller images scale down proportionally
                                                screen_width_px = 1920  # Typical screen width
                                                max_width_inches = 6.0
                                                calculated_width = min(max_width_inches, (width_px / screen_width_px) * max_width_inches)
                                                # Ensure minimum readable size
                                                final_width = max(1.5, calculated_width)
                                            
                                            table = doc.add_table(rows=1, cols=1)
                                            table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            cell = table.cell(0, 0)
                                            paragraph = cell.paragraphs[0]
                                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                            run = paragraph.add_run()
                                            run.add_picture(img_path, width=Inches(final_width))
                                            cap = cell.add_paragraph(bullet_text)
                                            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    except Exception:
                                        pass
                                continue
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Inches(0.05)
                    p.paragraph_format.line_spacing = 1.15
            # Subtasks as checkboxes
            if step.subtasks.exists():
                for idx, t in enumerate(step.subtasks.all().order_by('order', 'id')):
                    box = '☑' if t.completed else '☐'
                    doc.add_paragraph(f'{box} {t.text}')
                    
                    # Show images attached to this specific subtask
                    subtask_images = step.images.filter(subtask_index=idx)
                    if subtask_images.exists():
                        for jim in subtask_images.order_by('order', 'id'):
                            try:
                                img_path = os.path.join(settings.MEDIA_ROOT, str(jim.image))
                                if os.path.exists(img_path):
                                    # Get image dimensions to scale appropriately
                                    from PIL import Image
                                    with Image.open(img_path) as img:
                                        width_px, height_px = img.size
                                        # Calculate width based on image size relative to typical screen (1920px)
                                        # Full screen = 6 inches, smaller images scale down proportionally
                                        screen_width_px = 1920  # Typical screen width
                                        max_width_inches = 6.0
                                        calculated_width = min(max_width_inches, (width_px / screen_width_px) * max_width_inches)
                                        # Ensure minimum readable size
                                        final_width = max(1.5, calculated_width)
                                    
                                    p = doc.add_paragraph()
                                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    r = p.add_run()
                                    r.add_picture(img_path, width=Inches(final_width))
                                    # Add caption
                                    cap = doc.add_paragraph(f"Image for: {t.text[:50]}{'...' if len(t.text) > 50 else ''}")
                                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cap.paragraph_format.font.size = Inches(0.1)
                            except Exception:
                                pass

            # Include pasted job images (documentation)
            if step.images.exists():
                for jim in step.images.all().order_by('order', 'id'):
                    try:
                        img_path = os.path.join(settings.MEDIA_ROOT, str(jim.image))
                        if os.path.exists(img_path):
                            # Get image dimensions to scale appropriately
                            from PIL import Image
                            with Image.open(img_path) as img:
                                width_px, height_px = img.size
                                # Calculate width based on image size relative to typical screen (1920px)
                                # Full screen = 6 inches, smaller images scale down proportionally
                                screen_width_px = 1920  # Typical screen width
                                max_width_inches = 6.0
                                calculated_width = min(max_width_inches, (width_px / screen_width_px) * max_width_inches)
                                # Ensure minimum readable size
                                final_width = max(1.5, calculated_width)
                            
                            p = doc.add_paragraph()
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            r = p.add_run()
                            r.add_picture(img_path, width=Inches(final_width))
                    except Exception:
                        pass

            # Include step notes (emphasized title + boxed area)
            if step.notes:
                doc.add_heading('NOTES', level=3)
                # Box container using single-cell table with borders
                table = doc.add_table(rows=1, cols=1)
                cell = table.cell(0, 0)
                # Apply borders
                try:
                    from docx.oxml.shared import OxmlElement, qn
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right']:
                        border = OxmlElement(f'w:{border_name}')
                        border.set(qn('w:val'), 'single')
                        border.set(qn('w:sz'), '16')
                        border.set(qn('w:space'), '0')
                        border.set(qn('w:color'), '555555')
                        tcBorders.append(border)
                    tcPr.append(tcBorders)
                except Exception:
                    pass
                # Add notes content inside the box
                # Clear default empty paragraph
                cell.text = ''
                for ln in step.notes.split('\n'):
                    if ln.strip().startswith('- '):
                        p = cell.add_paragraph(ln.strip()[2:].strip(), style='List Bullet')
                    else:
                        p = cell.add_paragraph(ln)

            if show_attachments:
                src = source_by_order.get(step.order)
                if src and src.images.exists():
                    for img in src.images.all():
                        if img.substep_index is None:
                            try:
                                img_path = os.path.join(settings.MEDIA_ROOT, str(img.image))
                                if os.path.exists(img_path):
                                    # Get image dimensions to scale appropriately
                                    from PIL import Image
                                    with Image.open(img_path) as img_pil:
                                        width_px, height_px = img_pil.size
                                        # Calculate width based on image size relative to typical screen (1920px)
                                        # Full screen = 6 inches, smaller images scale down proportionally
                                        screen_width_px = 1920  # Typical screen width
                                        max_width_inches = 6.0
                                        calculated_width = min(max_width_inches, (width_px / screen_width_px) * max_width_inches)
                                        # Ensure minimum readable size
                                        final_width = max(1.5, calculated_width)
                                    
                                    table = doc.add_table(rows=1, cols=1)
                                    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    cell = table.cell(0, 0)
                                    paragraph = cell.paragraphs[0]
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    run = paragraph.add_run()
                                    run.add_picture(img_path, width=Inches(final_width))
                            except Exception:
                                pass
    doc_io = BytesIO()
    doc.settings.odd_and_even_pages_header_footer = False
    doc.save(doc_io)
    doc_io.seek(0)
    from datetime import datetime
    timestamp_date = datetime.now().strftime('%m-%d-%y')
    timestamp_time = datetime.now().strftime('%I-%M%p').lower()
    safe_title = ''.join(c for c in job.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    status_text = job.get_status_display().upper()
    filename = f"{safe_title} {status_text} -- {timestamp_date} -- {timestamp_time}.docx"
    response = HttpResponse(doc_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    # Force download to Downloads folder and auto-open
    response['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response['X-Content-Type-Options'] = 'nosniff'
    return response

@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def template_publish(request, tpl_id: int):
    template = get_object_or_404(ProcessTemplate, id=tpl_id)
    is_active = request.POST.get('is_active') in ['1', 'true', 'True', 'on']
    template.is_active = is_active
    template.save(update_fields=['is_active', 'updated_at'])
    return JsonResponse({'ok': True, 'is_active': template.is_active})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def template_create_from_selected(request):
    ids = request.POST.getlist('process_ids[]') or request.POST.getlist('process_ids')
    name = (request.POST.get('name') or '').strip()
    module_id = request.POST.get('module')
    if not ids or not name or not module_id:
        return JsonResponse({'ok': False, 'error': 'Name, module and processes are required'}, status=400)
    
    try:
        source_processes = Process.objects.filter(id__in=ids).order_by('order')
        module = Module.objects.get(id=module_id)
    except (Process.DoesNotExist, Module.DoesNotExist):
        return JsonResponse({'ok': False, 'error': 'Invalid process or module'}, status=400)
    
    if not source_processes.exists():
        return JsonResponse({'ok': False, 'error': 'No valid processes found'}, status=400)
    
    with transaction.atomic():
        # Create a new combined process from the selected processes
        combined_process = Process.objects.create(
            name=f"Template: {name}",
            module=module,
            description=f"Combined template from {len(source_processes)} processes",
            summary=f"Template combining: {', '.join([p.name for p in source_processes])}"
        )
        
        # Copy steps from all selected processes in order
        step_order = 1
        for process in source_processes:
            for step in process.steps.all().order_by('order'):
                # Create new step in combined process
                new_step = Step.objects.create(
                    process=combined_process,
                    order=step_order,
                    title=step.title,
                    details=step.details
                )
                step_order += 1
                
                # Copy step images
                for step_image in step.images.all():
                    StepImage.objects.create(
                        step=new_step,
                        image=step_image.image,
                        order=step_image.order,
                        substep_index=step_image.substep_index
                    )
        
        # Create template linked to the combined process
        template = ProcessTemplate.objects.create(
            name=name,
            module=module,
            source_process=combined_process,
            description=combined_process.summary or combined_process.description or '',
            is_active=True
        )
        
        # Auto-sync template from the combined process
        from .services.templates import sync_process_to_template
        sync_process_to_template(combined_process.id)
        
        return JsonResponse({'ok': True, 'template_id': template.id, 'process_id': combined_process.id})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_subtask_toggle(request, subtask_id: int):
    subtask = get_object_or_404(JobSubtask.objects.select_related('job_step', 'job_step__job'), id=subtask_id)
    completed = request.POST.get('completed') in ['1', 'true', 'True', 'on']
    subtask.completed = completed
    subtask.save(update_fields=['completed', 'updated_at'])
    # If all subtasks completed for a step, mark step completed; else if any started, mark in_progress
    step = subtask.job_step
    if step.subtasks.exists():
        if step.subtasks.filter(completed=False).count() == 0:
            _update_job_step(step.job, step, 'completed')
        else:
            _update_job_step(step.job, step, 'in_progress')
    # Update job status based on any completed subtasks overall
    job = step.job
    from django.db.models import Q
    any_done = JobSubtask.objects.filter(job_step__job=job, completed=True).exists()
    if any_done and job.status == 'not_started':
        job.status = 'in_progress'
        if not job.started_at:
            job.started_at = timezone.now()
        job.save(update_fields=['status', 'started_at', 'updated_at'])
    if not any_done:
        # If nothing done and no steps in progress/completed, revert to not_started
        if not job.steps.filter(Q(status='in_progress') | Q(status='completed')).exists():
            job.status = 'not_started'
            job.started_at = None
            job.save(update_fields=['status', 'started_at', 'updated_at'])
    return JsonResponse({'ok': True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_step_notes_update(request, job_step_id: int):
    step = get_object_or_404(JobStep, id=job_step_id)
    step.notes = request.POST.get('notes', '')
    step.save(update_fields=['notes', 'updated_at'])
    return JsonResponse({'ok': True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_step_image_upload(request, job_step_id: int):
    step = get_object_or_404(JobStep, id=job_step_id)
    img_file = request.FILES.get('image') or request.FILES.get('file')
    if not img_file:
        return JsonResponse({'ok': False, 'error': 'No image provided'}, status=400)
    sub_idx = request.POST.get('subtask_index')
    try:
        sub_idx = int(sub_idx) if sub_idx not in (None, '') else None
    except (TypeError, ValueError):
        sub_idx = None
    max_order = step.images.aggregate(models.Max('order')).get('order__max') or 0
    jsi = JobStepImage.objects.create(job_step=step, image=img_file, order=max_order + 1, subtask_index=sub_idx)
    return JsonResponse({'ok': True, 'id': jsi.id, 'url': jsi.image.url})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_step_image_delete(request, job_step_id: int, image_id: int):
    step = get_object_or_404(JobStep, id=job_step_id)
    img = get_object_or_404(JobStepImage, id=image_id, job_step=step)
    img.delete()
    return JsonResponse({'ok': True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_step_attachment_upload(request, job_step_id: int):
    """Upload file attachment to job step with auto-conversion to PDF"""
    from .models import JobStepAttachment
    from .services.file_conversion import convert_file_to_pdf, get_file_icon, format_file_size
    import mimetypes
    import os
    
    step = get_object_or_404(JobStep, id=job_step_id)
    
    if 'file' not in request.FILES:
        return JsonResponse({'ok': False, 'error': 'No file provided'}, status=400)
    
    file = request.FILES['file']
    subtask_index = request.POST.get('subtask_index')
    
    # Validate file size (25MB limit)
    max_size = 25 * 1024 * 1024  # 25MB
    if file.size > max_size:
        return JsonResponse({'ok': False, 'error': 'File too large. Maximum size is 25MB.'}, status=400)
    
    # Get file info
    original_name = file.name
    mime_type = mimetypes.guess_type(original_name)[0] or 'application/octet-stream'
    file_extension = original_name.split('.')[-1].lower() if '.' in original_name else ''
    
    # Create attachment
    attachment = JobStepAttachment.objects.create(
        job_step=step,
        file=file,
        original_name=original_name,
        mime_type=mime_type,
        file_size=file.size,
        subtask_index=int(subtask_index) if subtask_index else None,
        order=step.attachments.count() + 1
    )
    
    # Try to convert to PDF if supported
    if attachment.is_convertible:
        try:
            # Save file temporarily for conversion
            temp_path = os.path.join(settings.MEDIA_ROOT, 'temp', f'temp_{attachment.id}_{original_name}')
            os.makedirs(os.path.dirname(temp_path), exist_ok=True)
            
            with open(temp_path, 'wb') as f:
                for chunk in file.chunks():
                    f.write(chunk)
            
            # Convert to PDF
            pdf_path = convert_file_to_pdf(temp_path, file_extension)
            
            # Save PDF preview
            with open(pdf_path, 'rb') as pdf_file:
                attachment.pdf_preview.save(
                    f'preview_{attachment.id}.pdf',
                    File(pdf_file),
                    save=True
                )
            
            attachment.conversion_status = 'success'
            attachment.save()
            
            # Clean up temp files
            os.unlink(temp_path)
            os.unlink(pdf_path)
            
        except Exception as e:
            logger.error(f"File conversion failed for {original_name}: {str(e)}")
            attachment.conversion_status = 'failed'
            attachment.save()
            # Return error info for debugging
            return JsonResponse({
                'ok': True,
                'id': attachment.id,
                'url': attachment.file.url,
                'original_name': original_name,
                'file_size': format_file_size(file.size),
                'icon': get_file_icon(file_extension),
                'has_preview': False,
                'preview_url': None,
                'conversion_error': str(e)
            })
    
    return JsonResponse({
        'ok': True,
        'id': attachment.id,
        'url': attachment.file.url,
        'original_name': original_name,
        'file_size': format_file_size(file.size),
        'icon': get_file_icon(file_extension),
        'has_preview': attachment.has_preview,
        'preview_url': attachment.pdf_preview.url if attachment.has_preview else None
    })


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_step_attachment_delete(request, job_step_id: int, attachment_id: int):
    """Delete file attachment from job step"""
    from .models import JobStepAttachment
    
    step = get_object_or_404(JobStep, id=job_step_id)
    attachment = get_object_or_404(JobStepAttachment, id=attachment_id, job_step=step)
    attachment.delete()
    return JsonResponse({'ok': True})


@login_required
@require_app_access('process_creator', action='edit')
@require_POST
def job_update_from_template(request, job_id: int):
    job = get_object_or_404(Job.objects.select_related('template'), id=job_id)
    tpl = job.template
    # Merge template changes into job while retaining subtask completion by index
    with transaction.atomic():
        # Build maps of current job steps and subtasks
        job_steps_by_order = {js.order: js for js in job.steps.all().prefetch_related('subtasks')}
        # Ensure JobSteps exist for each TemplateStep
        for ts in tpl.steps.all().order_by('order', 'id'):
            js = job_steps_by_order.get(ts.order)
            if js is None:
                js = JobStep.objects.create(job=job, order=ts.order, title=ts.title, details=ts.details or '')
                job_steps_by_order[ts.order] = js
            else:
                js.title = ts.title
                js.details = ts.details or ''
                js.save(update_fields=['title', 'details', 'updated_at'])
            # Sync subtasks - preserve ALL existing completions
            # Extract bullets from TemplateStep details
            bullets = []
            if ts.details:
                for line in ts.details.split('\n'):
                    if line.strip().startswith('- '):
                        bullets.append(line.strip()[2:].strip())
            
            # Get existing subtasks and their completion status
            existing_subtasks = list(js.subtasks.all().order_by('order', 'id'))
            completed_texts = {st.text: st.completed for st in existing_subtasks}
            
            # Clear existing subtasks
            js.subtasks.all().delete()
            
            # Create new subtasks, preserving completion status by text matching
            for i, text in enumerate(bullets):
                completed = completed_texts.get(text, False)  # Keep completion if text matches
                JobSubtask.objects.create(
                    job_step=js, 
                    order=i, 
                    text=text, 
                    completed=completed
                )
        # Reorder any job steps to match template orders (preserve others)
        for js in job.steps.all():
            # if its order not in template, keep as is; else already aligned
            pass
        job.template_version_at_create = tpl.version
        job.save(update_fields=['template_version_at_create', 'updated_at'])
    return redirect('process_creator:job_detail', job_id=job.id)
