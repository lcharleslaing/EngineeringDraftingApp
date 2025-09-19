from django.shortcuts import get_object_or_404, redirect, render
from django.http import JsonResponse, HttpResponse
from django.views.decorators.http import require_POST
from django.db import transaction
from django.contrib.auth.decorators import login_required
from django.db import models
from django.template.loader import render_to_string
from django.conf import settings
from django.utils import timezone
from .models import Module, Process, Step, StepImage, StepLink, StepFile, AIInteraction
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json
import openai
import re
from decimal import Decimal
from io import BytesIO
from urllib.parse import quote
import subprocess
import tempfile


def markdown_to_plain_text(text):
    """Convert basic Markdown formatting to plain text with proper formatting"""
    if not text:
        return ""
    
    # Convert headers
    text = re.sub(r'^### (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'\1', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.+)$', r'\1', text, flags=re.MULTILINE)
    
    # Convert bold text
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    
    # Convert italic text
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    
    # Convert bullet points
    text = re.sub(r'^\* (.+)$', r'• \1', text, flags=re.MULTILINE)
    text = re.sub(r'^- (.+)$', r'• \1', text, flags=re.MULTILINE)
    
    # Convert numbered lists
    text = re.sub(r'^(\d+)\. (.+)$', r'\1. \2', text, flags=re.MULTILINE)
    
    # Clean up extra whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()


@login_required
def process_list(request):
    # Get all modules for the dropdown
    modules = Module.objects.all()
    
    # Get selected module from request
    selected_module_id = request.GET.get('module')
    selected_module = None
    
    if selected_module_id:
        try:
            selected_module = Module.objects.get(id=selected_module_id)
            processes = Process.objects.filter(module=selected_module).order_by('order', 'name')
        except Module.DoesNotExist:
            processes = Process.objects.all().order_by('order', 'name')
    else:
        processes = Process.objects.all().order_by('order', 'name')
    
    return render(request, "process_creator/list.html", {
        "processes": processes,
        "modules": modules,
        "selected_module": selected_module
    })


@login_required
@require_POST
def module_create(request):
    """Create a new Module via AJAX."""
    data = json.loads(request.body or '{}') if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    name = (data.get('name') or '').strip()
    description = (data.get('description') or '').strip()
    if not name:
        return JsonResponse({"ok": False, "error": "Name is required"}, status=400)
    # Prevent duplicates (case-insensitive)
    existing = Module.objects.filter(name__iexact=name).first()
    if existing:
        return JsonResponse({"ok": True, "id": existing.id, "name": existing.name})
    module = Module.objects.create(name=name, description=description)
    return JsonResponse({"ok": True, "id": module.id, "name": module.name})


@login_required
def process_create(request):
    if request.method == "POST":
        data = json.loads(request.body) if request.headers.get('Content-Type','').startswith('application/json') else request.POST
        name = (data.get('name') or '').strip()
        module_id = data.get('module')
        if not name:
            return JsonResponse({"ok": False, "error": "Name is required"}, status=400)
        module = None
        if module_id:
            try:
                module = Module.objects.get(id=module_id)
            except Module.DoesNotExist:
                pass
        process = Process.objects.create(name=name, module=module)
        return JsonResponse({"ok": True, "id": process.id, "url": f"/process-creator/{process.id}/"})
    return JsonResponse({"ok": False, "error": "Method not allowed"}, status=405)


@login_required
def process_edit(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    modules = Module.objects.all()
    return render(request, "process_creator/edit.html", {"process": process, "modules": modules})


@login_required
@require_POST
def process_update(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    data = json.loads(request.body) if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    for field in ['name', 'description', 'notes', 'summary', 'analysis', 'summary_instructions', 'analysis_instructions']:
        if field in data:
            setattr(process, field, data[field])
    if 'module' in data:
        module_id = data['module']
        if module_id:
            try:
                process.module = Module.objects.get(id=module_id)
            except Module.DoesNotExist:
                process.module = None
        else:
            process.module = None
    process.save()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def process_delete(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    process.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_add(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    data = json.loads(request.body) if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    title = (data.get('title') or 'New Step').strip()
    max_order = process.steps.aggregate(models.Max('order')).get('order__max') or 0
    step = Step.objects.create(process=process, title=title, order=max_order + 1)
    return JsonResponse({"ok": True, "id": step.id, "order": step.order})


@login_required
@require_POST
def step_update(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    data = json.loads(request.body) if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    for field in ['title', 'details']:
        if field in data:
            setattr(step, field, data[field])
    step.save()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_delete(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    step.delete()
    # Reorder remaining steps
    for i, remaining_step in enumerate(process.steps.all().order_by('id'), 1):
        remaining_step.order = i
        remaining_step.save()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_reorder(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    order_list = request.POST.getlist("order[]")
    with transaction.atomic():
        for index, step_id in enumerate(order_list, start=1):
            Step.objects.filter(process=process, id=step_id).update(order=index)
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_image_upload(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    file = request.FILES.get('file')
    if not file:
        return JsonResponse({"ok": False, "error": "No file uploaded"}, status=400)
    max_order = step.images.aggregate(models.Max('order')).get('order__max') or 0
    si = StepImage.objects.create(step=step, image=file, order=max_order + 1)
    return JsonResponse({"ok": True, "id": si.id, "url": si.image.url})


@login_required
@require_POST
def step_image_delete(request, pk: int, step_id: int, image_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    si = get_object_or_404(StepImage, pk=image_id, step=step)
    si.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_link_add(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    data = json.loads(request.body) if request.headers.get('Content-Type','').startswith('application/json') else request.POST
    title = (data.get('title') or '').strip()
    url = (data.get('url') or '').strip()
    if not title or not url:
        return JsonResponse({"ok": False, "error": "Title and URL are required"}, status=400)
    max_order = step.links.aggregate(models.Max('order')).get('order__max') or 0
    link = StepLink.objects.create(step=step, title=title, url=url, order=max_order + 1)
    return JsonResponse({"ok": True, "id": link.id, "title": link.title, "url": link.url})


@login_required
@require_POST
def step_link_delete(request, pk: int, step_id: int, link_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    link = get_object_or_404(StepLink, pk=link_id, step=step)
    link.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_file_upload(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    file = request.FILES.get('file')
    if not file:
        return JsonResponse({"ok": False, "error": "No file uploaded"}, status=400)

    original_filename = file.name
    file_extension = os.path.splitext(original_filename)[1].lower()
    
    # Determine if conversion is needed
    convert_to_pdf = False
    converter_path = None
    error_message = None

    if file_extension == '.pdf':
        pass # No conversion needed
    elif file_extension == '.dwg':
        converter_path = getattr(settings, 'ODA_CONVERTER_PDF', None)
        if not converter_path or not os.path.exists(converter_path):
            error_message = "DWG converter not configured or not found. Please set ODA_CONVERTER_PDF in settings"
        else:
            convert_to_pdf = True
    elif file_extension == '.idw':
        converter_path = getattr(settings, 'INVENTOR_IDW_TO_PDF', None)
        if not converter_path or not os.path.exists(converter_path):
            error_message = "IDW converter not configured or not found. Please set INVENTOR_IDW_TO_PDF in settings"
        else:
            convert_to_pdf = True
    else:
        # For other file types, just save the original file and create a link
        max_order = step.files.aggregate(models.Max('order')).get('order__max') or 0
        sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
        
        # Create a link for the uploaded file
        link_title = os.path.basename(original_filename)
        link_url = sf.file.url
        max_link_order = step.links.aggregate(models.Max('order')).get('order__max') or 0
        StepLink.objects.create(step=step, title=link_title, url=link_url, order=max_link_order + 1)

        return JsonResponse({
            "ok": True, 
            "id": sf.id, 
            "url": sf.file.url, 
            "name": os.path.basename(sf.file.name),
            "link_id": sf.links.first().id if sf.links.exists() else None,
            "link_title": link_title,
            "link_url": link_url,
        })

    if error_message:
        return JsonResponse({"ok": False, "error": error_message}, status=200)

    if convert_to_pdf:
        # Placeholder for actual conversion logic
        return JsonResponse({"ok": False, "error": f"Conversion for {file_extension} not yet fully implemented or converter failed."}, status=200)
    
    # If it's a PDF or conversion was successful, save it
    max_order = step.files.aggregate(models.Max('order')).get('order__max') or 0
    sf = StepFile.objects.create(step=step, file=file, order=max_order + 1)
    
    # Generate a human-readable timestamp for display
    timestamp_str = sf.uploaded_at.strftime("%m%d%y-%I%M%S%p").lower()
    base_name, ext = os.path.splitext(original_filename)
    display_name = f"{base_name} ({timestamp_str}){ext}"

    return JsonResponse({"ok": True, "id": sf.id, "url": sf.file.url, "name": display_name})


@login_required
@require_POST
def step_file_delete(request, pk: int, step_id: int, file_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    sf = get_object_or_404(StepFile, pk=file_id, step=step)
    sf.delete()
    return JsonResponse({"ok": True})


@login_required
@require_POST
def step_images_reorder(request, pk: int, step_id: int):
    process = get_object_or_404(Process, pk=pk)
    step = get_object_or_404(Step, pk=step_id, process=process)
    order_list = request.POST.getlist("order[]")
    with transaction.atomic():
        for index, img_id in enumerate(order_list, start=1):
            StepImage.objects.filter(step=step, id=img_id).update(order=index)
    return JsonResponse({"ok": True})


@login_required
def process_stats(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Calculate statistics
    step_count = process.steps.count()
    
    # Count substeps (lines starting with -)
    substep_count = 0
    for step in process.steps.all():
        if step.details:
            # Count lines that start with - (substeps)
            substep_count += len([line for line in step.details.split('\n') if line.strip().startswith('-')])
    
    total_steps = step_count + substep_count
    image_count = sum(step.images.count() for step in process.steps.all())
    
    # Calculate text lengths
    description_length = len(process.description) if process.description else 0
    notes_length = len(process.notes) if process.notes else 0
    summary_length = len(process.summary) if process.summary else 0
    analysis_length = len(process.analysis) if process.analysis else 0
    
    # Calculate total text length from all step details
    step_details_length = 0
    for step in process.steps.all():
        if step.details:
            step_details_length += len(step.details)
    
    # Total text length (all text content combined)
    total_text_length = description_length + notes_length + summary_length + analysis_length + step_details_length
    
    # Format dates
    from django.utils import timezone
    created_at = process.created_at.strftime('%B %d, %Y at %I:%M %p')
    updated_at = process.updated_at.strftime('%B %d, %Y at %I:%M %p')
    
    stats = {
        'created_at': created_at,
        'updated_at': updated_at,
        'step_count': step_count,
        'substep_count': substep_count,
        'total_steps': total_steps,
        'image_count': image_count,
        'description_length': description_length,
        'notes_length': notes_length,
        'summary_length': summary_length,
        'analysis_length': analysis_length,
        'step_details_length': step_details_length,
        'total_text_length': total_text_length,
    }
    
    return JsonResponse(stats)


def call_openai_api(prompt, model="gpt-4o-mini", max_tokens=2000):
    """Helper function to call OpenAI API and track usage"""
    try:
        if not getattr(settings, 'OPENAI_API_KEY', None):
            return {
                'success': False,
                'error': 'OpenAI API key not configured. Please set OPENAI_API_KEY in settings.',
                'tokens_used': 0,
                'cost': Decimal('0.00')
            }
        
        client = openai.OpenAI(api_key=settings.OPENAI_API_KEY)
        
        response = client.chat.completions.create(
            model=model,
            messages=[{"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.7
        )
        
        # Calculate cost (approximate for gpt-4o-mini)
        tokens_used = response.usage.total_tokens
        cost_per_1k_tokens = Decimal('0.00015')  # Approximate cost for gpt-4o-mini
        cost = (tokens_used / 1000) * cost_per_1k_tokens
        
        return {
            'success': True,
            'response': response.choices[0].message.content,
            'tokens_used': tokens_used,
            'cost': cost
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'tokens_used': 0,
            'cost': Decimal('0.00')
        }


@login_required
@require_POST
def process_summary(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Build the prompt from process steps
    steps_text = ""
    for step in process.steps.all():
        steps_text += f"Step {step.order}: {step.title}\n"
        if step.details:
            steps_text += f"{step.details}\n"
        steps_text += "\n"
    
    prompt = f"{process.summary_instructions}\n\nProcess Steps:\n{steps_text}"
    
    result = call_openai_api(prompt)
    
    if result['success']:
        process.summary = result['response']
        process.last_ai_update = timezone.now()
        process.save()
        
        # Log the interaction
        AIInteraction.objects.create(
            process=process,
            interaction_type='summary',
            prompt_sent=prompt,
            response_received=result['response'],
            tokens_used=result['tokens_used'],
            cost=result['cost']
        )
        
        return JsonResponse({
            "ok": True, 
            "summary": result['response'],
            "tokens_used": result['tokens_used'],
            "cost": float(result['cost'])
        })
    else:
        return JsonResponse({
            "ok": False, 
            "error": result['error']
        }, status=500)


@login_required
@require_POST
def process_analyze(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Build the prompt from process steps
    steps_text = ""
    for step in process.steps.all():
        steps_text += f"Step {step.order}: {step.title}\n"
        if step.details:
            steps_text += f"{step.details}\n"
        steps_text += "\n"
    
    prompt = process.analysis_instructions.replace("{PASTE PROCESS STEPS / NOTES / CONTEXT HERE}", steps_text)
    
    result = call_openai_api(prompt, max_tokens=4000)
    
    if result['success']:
        process.analysis = result['response']
        process.last_ai_update = timezone.now()
        process.save()
        
        # Log the interaction
        AIInteraction.objects.create(
            process=process,
            interaction_type='analysis',
            prompt_sent=prompt,
            response_received=result['response'],
            tokens_used=result['tokens_used'],
            cost=result['cost']
        )
        
        return JsonResponse({
            "ok": True, 
            "analysis": result['response'],
            "tokens_used": result['tokens_used'],
            "cost": float(result['cost'])
        })
    else:
        return JsonResponse({
            "ok": False, 
            "error": result['error']
        }, status=500)


@login_required
def process_pdf(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Render the template
    html_string = render_to_string('process_creator/print.html', {
        'process': process,
        'request': request
    })
    
    # Create PDF
    result = BytesIO()
    pdf = pisa.pisaDocument(BytesIO(html_string.encode("UTF-8")), result)
    
    if not pdf.err:
        # Generate filename with process title and timestamp
        from datetime import datetime
        timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
        safe_title = "".join(c for c in process.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_title = safe_title.replace(' ', '-')
        filename = f"Process-{safe_title}-{timestamp}.pdf"
        
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
        response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
        response['Pragma'] = 'no-cache'
        response['Expires'] = '0'
        return response
    else:
        return HttpResponse("Error generating PDF", status=500)


@login_required
def process_word(request, pk: int):
    process = get_object_or_404(Process, pk=pk)
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(process.name, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add description
    if process.description:
        doc.add_heading('Description', level=1)
        doc.add_paragraph(process.description)
    
    # Add steps
    doc.add_heading('Process Steps', level=1)
    for step in process.steps.all():
        doc.add_heading(f"{step.order}. {step.title}", level=2)
        if step.details:
            # Convert markdown to plain text for Word
            plain_text = markdown_to_plain_text(step.details)
            doc.add_paragraph(plain_text)
    
    # Add summary
    if process.summary:
        doc.add_heading('Summary', level=1)
        plain_text = markdown_to_plain_text(process.summary)
        doc.add_paragraph(plain_text)
    
    # Add analysis
    if process.analysis:
        doc.add_heading('Analysis', level=1)
        plain_text = markdown_to_plain_text(process.analysis)
        doc.add_paragraph(plain_text)
    
    # Add notes
    if process.notes:
        doc.add_heading('Notes', level=1)
        doc.add_paragraph(process.notes)
    
    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    
    # Generate filename with process title and timestamp
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    safe_title = "".join(c for c in process.name if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_title = safe_title.replace(' ', '-')
    filename = f"Process-{safe_title}-{timestamp}.docx"
    
    response = HttpResponse(doc_io.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{quote(filename)}"'
    response['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response['Pragma'] = 'no-cache'
    response['Expires'] = '0'
    return response
